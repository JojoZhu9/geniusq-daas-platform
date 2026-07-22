from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "智能问数优化实施计划书.md"
OUTPUT = ROOT / "docs" / "智能问数优化实施计划书.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D9E2EC") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_run_font(run, size=11, bold=False, color="1F2937", name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 16, "1F4E79", 16, 8),
        ("Heading 2", 13, "2563A6", 12, 6),
        ("Heading 3", 11.5, "1F4E79", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=60, after=8)
    run = p.add_run("GeniusQ DaaS 智能问数 Demo")
    set_run_font(run, size=24, bold=True, color="0B2545", name="Calibri")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=22)
    run = p.add_run("实施计划书")
    set_run_font(run, size=20, bold=True, color="2563A6", name="Calibri")

    meta = doc.add_table(rows=4, cols=2)
    meta.autofit = False
    widths = [Inches(1.6), Inches(4.4)]
    rows = [
        ("文档版本", "V2.0"),
        ("更新日期", "2026-07-20"),
        ("适用范围", "本地 Demo 演示、实习项目交付、后续平台化开发参考"),
        ("项目定位", "独立全栈原型，不直接修改或合并生产平台源码"),
    ]
    for row, values in zip(meta.rows, rows):
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.width = widths[idx]
            set_cell_border(cell)
            if idx == 0:
                set_cell_shading(cell, "F2F4F7")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(paragraph, after=0)
            run = paragraph.add_run(value)
            set_run_font(run, size=10.5, bold=(idx == 0), color="1F2937")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本文件同步当前项目功能、运行逻辑、技术架构、测试验收和后续扩展建议。")
    set_run_font(run, size=10.5, color="4B5563")
    doc.add_page_break()


def add_code_block(doc: Document, lines: list[str]) -> None:
    if not lines:
        return
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=6, line=1.0)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("334155")


def add_bullet(doc: Document, text: str, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p, after=4, line=1.12)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color="1F2937")


def add_plain_paragraph(doc: Document, text: str) -> None:
    if not text.strip():
        return
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color="1F2937")


def build_docx() -> None:
    md = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    in_code = False
    code_lines: list[str] = []
    skip_title = True

    for raw_line in md.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if skip_title and line.startswith("# "):
            skip_title = False
            continue

        if not line.strip():
            continue

        if line.startswith(">"):
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            paragraph = doc.add_paragraph(text, style=f"Heading {level}")
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            add_bullet(doc, numbered.group(1), numbered=True)
            continue

        if line.startswith("- "):
            add_bullet(doc, line[2:])
            continue

        add_plain_paragraph(doc, line)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("GeniusQ DaaS Intelligent Query Demo · Implementation Plan")
    set_run_font(run, size=8.5, color="64748B")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
