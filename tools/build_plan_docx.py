from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BRAND = "GeniusQ DaaS Platform"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2EC")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def strip_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return text.strip()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 14, 7),
        ("Heading 2", 13, "2E74B5", 10, 5),
        ("Heading 3", 11.5, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")


def add_meta(doc: Document, line: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(line)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("667085")


def add_markdown_table(doc: Document, rows: list[str]) -> None:
    parsed = [[strip_inline(cell) for cell in row.strip().strip("|").split("|")] for row in rows]
    if len(parsed) < 2:
        return
    header = parsed[0]
    body = [row for row in parsed[2:] if len(row) == len(header)]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for idx, value in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = value
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row_values in body:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    col_count = len(header)
    if col_count == 2:
        widths = [2600, 6760]
    elif col_count == 3:
        widths = [1900, 3400, 4060]
    elif col_count == 4:
        widths = [1200, 2300, 3300, 2560]
    else:
        widths = [int(9360 / col_count)] * col_count
    set_table_width(table, widths)
    set_table_borders(table)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("344054")
    doc.add_paragraph()


def build_docx(md_path: Path, out_path: Path) -> None:
    doc = Document()
    configure_styles(doc)
    lines = md_path.read_text(encoding="utf-8-sig").splitlines()
    title = strip_inline(lines[0].lstrip("# ")) if lines and lines[0].startswith("# ") else f"{BRAND} 实施计划书"
    add_title(doc, title)

    index = 1
    meta_lines: list[str] = []
    while index < len(lines) and (not lines[index].strip() or lines[index].startswith(">")):
        if lines[index].startswith(">"):
            meta_lines.append(strip_inline(lines[index].lstrip("> ")))
        index += 1
    if meta_lines:
        add_meta(doc, " | ".join(meta_lines))

    table_buffer: list[str] = []
    code_buffer: list[str] = []
    in_code = False

    def flush_table() -> None:
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    def flush_code() -> None:
        nonlocal code_buffer
        if code_buffer:
            add_code_block(doc, code_buffer)
            code_buffer = []

    for raw_line in lines[index:]:
        line = raw_line.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_table()
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue
        if line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(strip_inline(line[2:]), level=1)
        elif line.startswith("## "):
            doc.add_heading(strip_inline(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(strip_inline(line[4:]), level=2)
        elif line.startswith("#### "):
            doc.add_heading(strip_inline(line[5:]), level=3)
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            p.add_run(strip_inline(re.sub(r"^\d+\.\s+", "", line)))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(strip_inline(line[2:]))
        else:
            doc.add_paragraph(strip_inline(line))

    flush_table()
    flush_code()

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{BRAND} Intelligent Query Demo")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("667085")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


if __name__ == "__main__":
    build_docx(Path(sys.argv[1]), Path(sys.argv[2]))
